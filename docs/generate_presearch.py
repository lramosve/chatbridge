from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

# ── Title Page ──
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('ChatBridge')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Pre-Search Planning Document')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = desc.add_run('Building an AI Chat Platform with Third-Party App Integration')
run.font.size = Pt(12)
run.font.italic = True

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Date: March 30, 2026\n').font.size = Pt(11)
meta.add_run('Based on: Chatbox (chatboxai/chatbox) fork\n').font.size = Pt(11)
meta.add_run('Sprint: 1 week (MVP due Tuesday)').font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# CASE STUDY ANALYSIS
# ══════════════════════════════════════════════════════════════
doc.add_heading('Case Study Analysis', level=1)

case_study_text = """\
TutorMeAI sits at a fascinating inflection point. They have built a defensible position in K-12 education \
not through superior AI, but through configurability: teachers shape the chatbot to fit their classroom. \
With 10,000 districts and 200,000 daily users, they are no longer a startup experiment; they are \
infrastructure. The question is no longer whether the product works, but what it should become next. \
Their answer is orchestration: letting third-party apps live inside the chat experience so students can \
play a chess game, work through flashcards, or interact with a physics simulator without ever leaving \
the conversation window.

The central engineering problem is deceptively simple to state and genuinely hard to solve. Third-party \
developers control their own UI and define their own tools. The platform cannot predict what any given \
app will do, yet it must render that app's interface inside the chat, pass it the right data, know when \
it is finished, and keep the conversation coherent throughout. A student says "let's play chess," a board \
appears, they ask mid-game "what should I do here?", the chatbot reasons about the board state and \
responds, the game ends, and the conversation continues naturally. Every handoff in that sequence is a \
potential failure point.

Two categories of trade-off dominate the design space. The first is trust and safety. These are children. \
A malicious or poorly built app could expose student data, display inappropriate content, or simply break \
in a way that confuses a ten-year-old. We address this through iframe sandboxing with strict Content \
Security Policy headers, ensuring third-party code cannot access the parent DOM, read cookies, or \
navigate the top frame. Apps run in an isolated context and communicate only through a structured \
postMessage protocol. This imposes friction on developers but makes the safety boundary enforceable \
rather than aspirational. The ethical obligation here is non-negotiable: convenience for developers must \
never come at the cost of student safety.

The second category is communication and state. The chatbot must stay aware of what is happening \
inside an app it does not control. We solve this by requiring apps to register a tool schema on the \
platform, a machine-readable contract that declares what the app can do, what parameters each action \
takes, and what it returns. The chatbot's LLM uses function calling to invoke these tools, and the app \
signals completion back through the same postMessage channel. App state is owned entirely by the app; \
the platform stores only a lightweight summary that gets injected into the conversation context so the \
LLM can reason about what happened.

We landed on building atop the existing Chatbox desktop client, which already provides multi-provider \
LLM support through the Vercel AI SDK, streaming responses, conversation history via IndexedDB, and \
an MCP-based tool system. Rather than rebuilding these capabilities, we extend them: the plugin \
registry becomes a new layer that translates third-party tool schemas into MCP-compatible tool \
definitions the AI SDK already knows how to call. This lets us ship a working integration faster and \
with fewer moving parts.

The hardest trade-off we accepted is scope. A production platform would need app review pipelines, \
automated content scanning, and granular teacher-level permissions for which apps are visible to which \
students. Within a one-week sprint, we cannot build all of that honestly. Instead, we focus on getting \
the core lifecycle right for three apps: one complex stateful app (chess), one public API integration \
(weather), and one requiring OAuth authentication (Spotify). If the handshake between chat and app is \
rock-solid for these three patterns, the safety and governance layers become straightforward to add \
later. A platform that works reliably for three apps is worth more than one that works intermittently \
for ten.\
"""

doc.add_paragraph(case_study_text)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PHASE 1: DEFINE YOUR CONSTRAINTS
# ══════════════════════════════════════════════════════════════
doc.add_heading('Phase 1: Define Your Constraints', level=1)

# --- 1. Scale & Load Profile ---
doc.add_heading('1. Scale & Load Profile', level=2)

doc.add_heading('Users at launch? In 6 months?', level=3)
doc.add_paragraph(
    'At launch (demo/grading context): 5-20 concurrent users. '
    'In 6 months (hypothetical TutorMeAI production rollout): 10,000-50,000 DAU. '
    'The architecture should not preclude horizontal scaling, but the MVP targets single-server performance.'
)

doc.add_heading('Traffic pattern: steady, spiky, or unpredictable?', level=3)
doc.add_paragraph(
    'Spiky. K-12 usage follows school schedules: peaks at 8-9 AM and 1-2 PM local time on weekdays, '
    'near-zero on weekends and holidays. Each spike lasts 45-90 minutes (class periods). '
    'This favors a serverless or auto-scaling backend over fixed-capacity infrastructure.'
)

doc.add_heading('How many concurrent app sessions per user?', level=3)
doc.add_paragraph(
    'One active app at a time per conversation, but a user may have multiple conversations open. '
    'Design for 1 active app session per conversation, up to 3 conversations per user. '
    'Apps that are not in the active viewport should be suspended (iframe unloaded) to conserve resources.'
)

doc.add_heading('Cold start tolerance for app loading?', level=3)
doc.add_paragraph(
    'Under 3 seconds is acceptable for initial app load. Users will see a skeleton/spinner. '
    'Subsequent interactions within a loaded app should feel instant (<200ms round-trip for postMessage). '
    'If an app exceeds 10 seconds on cold start, the platform shows a timeout message and offers retry.'
)

# --- 2. Budget & Cost Ceiling ---
doc.add_heading('2. Budget & Cost Ceiling', level=2)

doc.add_heading('Monthly spend limit?', level=3)
doc.add_paragraph(
    'MVP/development phase: $50-100/month total (LLM API costs dominate). '
    'Production projection at 10K users: $2,000-5,000/month depending on invocation frequency. '
    'Infrastructure (Vercel/Railway free tiers) should remain under $20/month for MVP.'
)

doc.add_heading('Pay-per-use acceptable or need fixed costs?', level=3)
doc.add_paragraph(
    'Pay-per-use is preferred for the MVP since traffic is low and unpredictable. '
    'At production scale, a reserved/committed-use discount from the LLM provider would reduce cost. '
    'Infrastructure should remain pay-per-use (serverless) to handle the spiky K-12 traffic pattern.'
)

doc.add_heading('LLM cost per tool invocation acceptable range?', level=3)
doc.add_paragraph(
    'Target: $0.005-0.02 per tool invocation (including the LLM reasoning step that decides to call the tool). '
    'A typical chess game might involve 5-10 tool calls, totaling $0.05-0.20 per game session. '
    'Using Claude Haiku or GPT-4o-mini for tool routing keeps per-invocation costs at the low end. '
    'Complex reasoning (e.g., chess move analysis) can use a more capable model selectively.'
)

doc.add_heading('Where will you trade money for time?', level=3)
doc.add_paragraph(
    'Managed services over self-hosted: Supabase for database (free tier), Vercel for deployment (free tier), '
    'Clerk or Supabase Auth for authentication (avoid building custom JWT infrastructure). '
    'Pre-built chess libraries (chess.js, chessboard.js) over implementing game logic from scratch. '
    'Vercel AI SDK (already in the codebase) over raw API calls to LLM providers.'
)

# --- 3. Time to Ship ---
doc.add_heading('3. Time to Ship', level=2)

doc.add_heading('MVP timeline?', level=3)
doc.add_paragraph(
    'MVP (basic chat + one working third-party app): 24 hours (Tuesday deadline). '
    'Early submission (full plugin system + 3 apps): 4 days (Friday). '
    'Final (polish, auth flows, documentation, deployment): 7 days (Sunday).'
)

doc.add_heading('Speed-to-market vs. long-term maintainability priority?', level=3)
doc.add_paragraph(
    'Speed-to-market for the sprint. The Chatbox fork gives us a mature chat UI, multi-provider LLM support, '
    'and streaming out of the box. We extend rather than rebuild. However, the plugin interface contract '
    '(tool schema format, postMessage protocol, registration API) must be designed carefully because '
    'third-party developers will build against it. A sloppy contract now creates painful migrations later. '
    'Decision: move fast on UI and integration plumbing, move carefully on the public API surface.'
)

doc.add_heading('Iteration cadence after launch?', level=3)
doc.add_paragraph(
    'Weekly during the sprint. Post-sprint (hypothetical production): biweekly releases for platform, '
    'continuous deployment for individual third-party apps (they control their own release cycle). '
    'The plugin contract should be versioned from day one (v1) so future changes do not break existing apps.'
)

# --- 4. Security & Sandboxing ---
doc.add_heading('4. Security & Sandboxing', level=2)

doc.add_heading('How will you isolate third-party app code?', level=3)
doc.add_paragraph(
    'Iframe-based sandboxing. Each third-party app renders in an <iframe> with the sandbox attribute: '
    'sandbox="allow-scripts allow-forms". Crucially, allow-same-origin is NOT set, preventing the app '
    'from accessing the parent window\'s cookies, localStorage, or DOM. Communication happens exclusively '
    'through window.postMessage with origin validation on both sides.'
)

doc.add_heading('What happens if a malicious app is registered?', level=3)
doc.add_paragraph(
    'Defense in depth: (1) Apps must be registered through an admin-approved process (not self-service in MVP). '
    '(2) The iframe sandbox prevents DOM access and cookie theft. (3) CSP headers restrict what the iframe can load. '
    '(4) The postMessage protocol only accepts known message types; unexpected messages are logged and dropped. '
    '(5) Rate limiting on tool invocations prevents resource exhaustion. '
    'A malicious app can at worst render inappropriate content inside its own iframe; it cannot exfiltrate data '
    'or affect other parts of the platform.'
)

doc.add_heading('Content Security Policy requirements?', level=3)
doc.add_paragraph(
    'Platform CSP: default-src \'self\'; frame-src https://<approved-app-domains>; script-src \'self\'. '
    'Iframe CSP (set by the app\'s own server): should be permissive enough for the app to function but '
    'the platform does not rely on the app\'s CSP for security (the sandbox attribute is the enforcement boundary). '
    'frame-ancestors directive on the platform prevents clickjacking.'
)

doc.add_heading('Data privacy between apps and chat context?', level=3)
doc.add_paragraph(
    'Apps receive only the data explicitly passed to them via tool invocations (structured parameters). '
    'They do not receive the full conversation history, user profile, or other apps\' state. '
    'The platform injects a summary of app results into the LLM context, not raw app data. '
    'Student PII (name, school) is never forwarded to third-party apps in the MVP.'
)

# --- 5. Team & Skill Constraints ---
doc.add_heading('5. Team & Skill Constraints', level=2)

doc.add_heading('Solo or team?', level=3)
doc.add_paragraph('Solo developer for this sprint.')

doc.add_heading('Languages/frameworks you know well?', level=3)
doc.add_paragraph(
    'Strong: TypeScript, React, Node.js, Python. '
    'Familiar: Electron, Vite, Zustand, Tailwind CSS. '
    'The Chatbox codebase uses all of these, which reduces ramp-up time significantly.'
)

doc.add_heading('Experience with iframe/postMessage communication?', level=3)
doc.add_paragraph(
    'Moderate. Have used postMessage for cross-origin widget embedding before. '
    'The key complexity here is building a reliable bidirectional protocol with proper origin validation, '
    'message typing, and error handling. Will reference the Channel Messaging API as a more structured '
    'alternative to raw postMessage if complexity warrants it.'
)

doc.add_heading('Familiarity with OAuth2 flows?', level=3)
doc.add_paragraph(
    'Familiar with the authorization code flow and refresh token handling. '
    'For the Spotify integration, we will use the Authorization Code with PKCE flow (recommended for '
    'public clients). The platform backend acts as the OAuth client; the iframe app never sees the raw tokens. '
    'Will use a library (e.g., Supabase Auth with Spotify provider, or simple-oauth2) rather than implementing '
    'the flow from scratch.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PHASE 2: ARCHITECTURE DISCOVERY
# ══════════════════════════════════════════════════════════════
doc.add_heading('Phase 2: Architecture Discovery', level=1)

# --- 6. Plugin Architecture ---
doc.add_heading('6. Plugin Architecture', level=2)

doc.add_heading('Iframe-based vs web component vs server-side rendering?', level=3)
doc.add_paragraph(
    'Decision: Iframe-based.\n\n'
    'Trade-off analysis:\n'
    '- Iframes: Best isolation (separate browsing context), supports any frontend framework, well-understood '
    'security model. Downside: cross-origin communication is asynchronous and requires a protocol; styling '
    'cannot be shared with the parent.\n'
    '- Web Components: Better styling integration, lighter weight. Downside: runs in the same JS context as '
    'the platform (no true isolation); a broken component can crash the host page. Unacceptable for '
    'third-party code in a children\'s product.\n'
    '- Server-side rendering: App UI is rendered on the server and streamed as HTML. Downside: no client-side '
    'interactivity without hydration, which brings us back to needing client-side code execution.\n\n'
    'Iframes win because isolation is non-negotiable for a K-12 platform running untrusted third-party code.'
)

doc.add_heading('How will apps register their tool schemas?', level=3)
doc.add_paragraph(
    'Apps provide a manifest file (JSON) at a well-known URL: <app-base-url>/chatbridge-manifest.json. '
    'The manifest declares:\n'
    '- App metadata (name, description, icon, version)\n'
    '- Tool definitions (name, description, parameters as JSON Schema, return type)\n'
    '- UI entry point URL (the iframe src)\n'
    '- Required auth type (none, api_key, oauth2)\n'
    '- Permissions requested (e.g., "read_conversation_context")\n\n'
    'The platform fetches and validates this manifest at registration time. Tool schemas are converted '
    'to MCP-compatible tool definitions that the Vercel AI SDK can call via function calling.'
)

doc.add_heading('Message passing protocol (postMessage, custom events, WebSocket)?', level=3)
doc.add_paragraph(
    'Decision: postMessage with a typed protocol.\n\n'
    'The protocol defines message types:\n'
    '- TOOL_INVOKE: Platform -> App (tool name + parameters)\n'
    '- TOOL_RESULT: App -> Platform (result data + status)\n'
    '- STATE_UPDATE: App -> Platform (lightweight state summary for LLM context)\n'
    '- APP_READY: App -> Platform (signals iframe has loaded)\n'
    '- APP_COMPLETE: App -> Platform (signals the app interaction is finished)\n'
    '- APP_ERROR: App -> Platform (error reporting)\n\n'
    'Each message includes a correlation ID for request-response matching and a timestamp for ordering. '
    'Origin validation is enforced on both sides. Messages that do not conform to the schema are dropped.'
)

doc.add_heading('How does the chatbot discover available tools at runtime?', level=3)
doc.add_paragraph(
    'At conversation start, the platform loads all registered app manifests and converts their tool schemas '
    'into the tool format expected by the Vercel AI SDK\'s streamText function. These tools are injected '
    'alongside any existing MCP tools (web search, knowledge base). When the user sends a message, the LLM '
    'sees all available tools and decides whether to invoke one based on the user\'s intent. '
    'Tool descriptions in the manifest should be written to help the LLM route correctly (e.g., '
    '"Start an interactive chess game in the chat" rather than "chess").'
)

# --- 7. LLM & Function Calling ---
doc.add_heading('7. LLM & Function Calling', level=2)

doc.add_heading('Which LLM provider for function calling?', level=3)
doc.add_paragraph(
    'Primary: Anthropic Claude (claude-sonnet-4-20250514) via the Vercel AI SDK\'s @ai-sdk/anthropic provider. '
    'Claude has strong function calling support and the Chatbox codebase already integrates it. '
    'Fallback: OpenAI GPT-4o via @ai-sdk/openai (also already integrated). '
    'The Vercel AI SDK abstracts the provider, so switching is a config change, not a code change.'
)

doc.add_heading('How will dynamic tool schemas be injected into the system prompt?', level=3)
doc.add_paragraph(
    'The Vercel AI SDK\'s streamText() accepts a tools parameter as a Record<string, tool>. '
    'At generation time, we merge three tool sources:\n'
    '1. Built-in tools (web search, knowledge base) - already in the codebase\n'
    '2. MCP tools (from connected MCP servers) - already in the codebase\n'
    '3. Plugin tools (from registered third-party app manifests) - new\n\n'
    'Plugin tools are dynamically generated from manifest schemas using the AI SDK\'s tool() helper. '
    'The tool\'s execute function sends a postMessage to the app\'s iframe and awaits the TOOL_RESULT response.'
)

doc.add_heading('Context window management with multiple app schemas?', level=3)
doc.add_paragraph(
    'Each tool definition consumes roughly 100-300 tokens of context. With 3 apps averaging 3 tools each, '
    'that is ~900-2,700 tokens for tool schemas, well within budget. At scale (20+ apps), we would need to '
    'prune: only inject tools for apps the user has explicitly enabled or that are relevant to the current '
    'conversation topic. For the MVP with 3 apps, all tools are always available. '
    'The existing context compaction system in Chatbox will summarize long conversations to keep total context '
    'under the model\'s window limit.'
)

doc.add_heading('Streaming responses while waiting for tool results?', level=3)
doc.add_paragraph(
    'The Vercel AI SDK supports streaming with tool calls natively. When the LLM decides to call a tool, '
    'the stream pauses, the tool executes (postMessage to iframe, await result), and the stream resumes '
    'with the LLM\'s response incorporating the tool result. During the tool execution pause, the UI shows '
    'a "Running [app name]..." indicator. If the tool call takes more than 30 seconds, a timeout error is '
    'returned to the LLM, which can then inform the user.'
)

# --- 8. Real-Time Communication ---
doc.add_heading('8. Real-Time Communication', level=2)

doc.add_heading('WebSocket vs SSE vs polling for chat?', level=3)
doc.add_paragraph(
    'The existing Chatbox architecture uses direct LLM API streaming (SSE from the AI provider) rather than '
    'a separate WebSocket layer. Since this is a desktop/web app making direct API calls (not routing through '
    'a custom backend for chat), we keep this architecture. The Vercel AI SDK handles streaming natively. '
    'No additional WebSocket server is needed for the MVP. If we add a shared backend later (e.g., for '
    'multi-device sync), WebSockets would be the right choice then.'
)

doc.add_heading('Separate channel for app-to-platform communication?', level=3)
doc.add_paragraph(
    'Yes. Chat communication (user <-> LLM) flows through the AI SDK\'s streaming API. '
    'App communication (platform <-> iframe) flows through postMessage. These are deliberately separate '
    'channels to avoid coupling. The platform\'s plugin controller acts as the bridge: it translates '
    'LLM tool calls into postMessage invocations and translates postMessage results back into tool responses '
    'for the AI SDK.'
)

doc.add_heading('How do you handle bidirectional state updates?', level=3)
doc.add_paragraph(
    'App -> Platform: The app sends STATE_UPDATE messages whenever its internal state changes meaningfully '
    '(e.g., a chess move is made). The platform stores the latest state summary in a per-conversation '
    'plugin state store (Zustand). This summary is injected into the next LLM call as a system message.\n\n'
    'Platform -> App: When the LLM invokes a tool, the platform sends a TOOL_INVOKE message to the app. '
    'The app processes it and sends back a TOOL_RESULT. The platform does not push unsolicited state to the app; '
    'the app is the source of truth for its own state.'
)

doc.add_heading('Reconnection and message ordering guarantees?', level=3)
doc.add_paragraph(
    'PostMessage within a single page session is synchronous and ordered (same-thread delivery). '
    'If the iframe reloads (navigation, crash), the platform detects the loss of the APP_READY heartbeat '
    'and prompts the user to reload the app. Message ordering is guaranteed by the browser\'s event loop. '
    'Each message includes a monotonic sequence number for debugging, but reordering is not expected in practice.'
)

# --- 9. State Management ---
doc.add_heading('9. State Management', level=2)

doc.add_heading('Where does chat state live? App state? Session state?', level=3)
doc.add_paragraph(
    'Chat state: Zustand store (chatStore.ts) persisted to IndexedDB via localforage. This is the existing '
    'Chatbox architecture and we keep it.\n\n'
    'App state: Owned by the app inside its iframe. The platform only stores a lightweight JSON summary '
    '(last STATE_UPDATE) in a new Zustand slice (pluginStore.ts), keyed by conversation ID + app ID.\n\n'
    'Session state: The combination of chat messages + plugin state summaries. When the LLM is called, '
    'active plugin state is appended to the system message so the LLM has context about what apps are doing.'
)

doc.add_heading('How do you merge app context back into conversation history?', level=3)
doc.add_paragraph(
    'When an app signals completion (APP_COMPLETE), the platform creates a synthetic assistant message: '
    '"[Chess app] Game completed. White won by checkmate in 24 moves." This message is added to the '
    'conversation history like any other message, making it visible to the user and available in the LLM\'s '
    'context for future turns. Tool invocation results are also stored as tool-role messages in the '
    'existing message format.'
)

doc.add_heading('State persistence across page refreshes?', level=3)
doc.add_paragraph(
    'Chat state persists (IndexedDB). Plugin state summaries persist (stored alongside chat state). '
    'App internal state: depends on the app. The platform sends the last known state summary to the app '
    'via a RESTORE_STATE message when the iframe reloads, but the app is responsible for using it. '
    'Apps that need durable state should use their own backend storage.'
)

doc.add_heading('What happens to the app state if the user closes the chat?', level=3)
doc.add_paragraph(
    'The iframe is destroyed. The platform\'s plugin state summary is retained in IndexedDB. '
    'When the user reopens the conversation, the app can be relaunched and the platform sends the last '
    'state summary. However, if the app does not support state restoration, the interaction starts fresh. '
    'The conversation history still contains the record of what happened, so the LLM has context.'
)

# --- 10. Authentication Architecture ---
doc.add_heading('10. Authentication Architecture', level=2)

doc.add_heading('Platform auth vs per-app auth?', level=3)
doc.add_paragraph(
    'Both, layered. Platform auth (Supabase Auth or Clerk) handles user identity for the ChatBridge '
    'platform itself. Per-app auth (OAuth2) handles authorization for external services like Spotify. '
    'The platform auth token is never shared with third-party apps. '
    'Per-app OAuth tokens are stored server-side and passed to the app via secure TOOL_INVOKE parameters.'
)

doc.add_heading('Token storage and refresh strategy?', level=3)
doc.add_paragraph(
    'Platform tokens: stored in Zustand + localStorage (existing Chatbox pattern). '
    'OAuth tokens for third-party services: stored in the backend database (Supabase), encrypted at rest. '
    'Refresh tokens are rotated automatically using the OAuth provider\'s refresh endpoint. '
    'The frontend never sees raw OAuth tokens; it receives a session-scoped proxy token from the platform backend.'
)

doc.add_heading('OAuth redirect handling within iframe context?', level=3)
doc.add_paragraph(
    'OAuth flows cannot happen inside a sandboxed iframe (redirects are blocked). '
    'Solution: When an app requires OAuth, the platform opens a popup window for the OAuth flow. '
    'The popup completes the authorization, the callback hits the platform backend, tokens are stored, '
    'and the popup closes. The platform then notifies the iframe that auth is complete via postMessage. '
    'This is the standard pattern used by Figma, Notion, and other platforms with embedded OAuth.'
)

doc.add_heading('How do you surface auth requirements to the user naturally?', level=3)
doc.add_paragraph(
    'When the LLM invokes a tool for an app that requires auth and the user has not yet authorized, '
    'the tool returns a special AUTH_REQUIRED status instead of a result. The LLM receives this and '
    'generates a natural-language prompt: "To use Spotify, I\'ll need you to connect your account. '
    'Click the button below to authorize." The platform renders an auth button inline in the chat. '
    'After auth completes, the original tool call is retried automatically.'
)

# --- 11. Database & Persistence ---
doc.add_heading('11. Database & Persistence', level=2)

doc.add_heading('Schema design for conversations, app registrations, sessions?', level=3)
doc.add_paragraph(
    'Conversations and messages: Keep the existing IndexedDB/localforage schema from Chatbox. '
    'No migration needed.\n\n'
    'App registrations (new): A JSON config file or lightweight SQLite table (via the existing @libsql/client) '
    'storing registered app manifests. Fields: app_id, name, manifest_url, manifest_cache, status '
    '(active/disabled), registered_at.\n\n'
    'Plugin sessions (new): Keyed by conversation_id + app_id. Fields: last_state_summary (JSON), '
    'tool_invocation_count, last_active_at.\n\n'
    'OAuth tokens (new, backend only): user_id, app_id, access_token (encrypted), refresh_token (encrypted), '
    'expires_at, scopes.'
)

doc.add_heading('How do you store tool invocation history?', level=3)
doc.add_paragraph(
    'Tool invocations are stored as tool-role messages in the conversation history (the existing message format '
    'supports role: "tool" with contentParts). This gives us invocation history for free: it is the '
    'conversation itself. For analytics (invocation counts, error rates), we add lightweight counters '
    'to the plugin session record.'
)

doc.add_heading('Read/write patterns and indexing strategy?', level=3)
doc.add_paragraph(
    'Read-heavy: conversation history is read on every page load and every LLM call. Writes happen on each '
    'new message. The existing IndexedDB setup handles this well for single-user desktop use. '
    'For the backend (app registrations, OAuth tokens), Supabase provides automatic indexing on primary keys. '
    'We add an index on (user_id, app_id) for the OAuth tokens table.'
)

doc.add_heading('Backup and disaster recovery?', level=3)
doc.add_paragraph(
    'MVP: Local data only (Chatbox\'s existing model). Conversation export to JSON is already supported. '
    'Production: Supabase provides automatic daily backups for PostgreSQL data. '
    'User-facing: add a "sync conversations" feature using Supabase Realtime + PostgreSQL in a future iteration.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# PHASE 3: POST-STACK REFINEMENT
# ══════════════════════════════════════════════════════════════
doc.add_heading('Phase 3: Post-Stack Refinement', level=1)

# --- 12. Security & Sandboxing Deep Dive ---
doc.add_heading('12. Security & Sandboxing Deep Dive', level=2)

doc.add_heading('Iframe sandbox attributes (allow-scripts, allow-same-origin)?', level=3)
doc.add_paragraph(
    'sandbox="allow-scripts allow-forms allow-popups"\n\n'
    '- allow-scripts: Required for apps to function (they are interactive).\n'
    '- allow-forms: Required for apps with form inputs (e.g., search bars).\n'
    '- allow-popups: Required for OAuth flows that open a popup.\n'
    '- allow-same-origin: NOT included. This is the critical security decision. Without it, the iframe '
    'cannot access the parent\'s cookies, localStorage, or session storage. The trade-off is that the app '
    'cannot use its own cookies either, so apps needing persistent state must use postMessage to ask the '
    'platform to store data on their behalf, or use their own backend.'
)

doc.add_heading('CSP headers for embedded content?', level=3)
doc.add_paragraph(
    'The platform sets: Content-Security-Policy: frame-src https://<allowlisted-domains>; '
    'frame-ancestors \'self\'. This restricts which domains can be embedded as iframes and prevents '
    'the platform itself from being embedded by malicious sites. '
    'Each registered app\'s domain is added to the frame-src allowlist at registration time.'
)

doc.add_heading('Preventing apps from accessing parent DOM?', level=3)
doc.add_paragraph(
    'The sandbox attribute without allow-same-origin enforces this at the browser level. '
    'Even if the app runs malicious JavaScript, it cannot call parent.document, window.top, or '
    'window.parent.postMessage with a spoofed origin. The browser\'s Same-Origin Policy is the enforcement '
    'mechanism. We do not rely on app developers being trustworthy.'
)

doc.add_heading('Rate limiting per app and per user?', level=3)
doc.add_paragraph(
    'Platform-side rate limiting on postMessage processing: max 100 messages/second per iframe. '
    'Tool invocations: max 10 per minute per app per user (prevents runaway loops). '
    'LLM API calls: governed by the existing Chatbox rate limiting + provider-side limits. '
    'If an app exceeds the postMessage rate limit, messages are dropped and an APP_ERROR is logged.'
)

# --- 13. Error Handling & Resilience ---
doc.add_heading('13. Error Handling & Resilience', level=2)

doc.add_heading('What happens when an app\'s iframe fails to load?', level=3)
doc.add_paragraph(
    'The platform sets a 10-second timeout after setting the iframe src. If no APP_READY message is '
    'received, the platform shows an error state in the chat: "The chess app failed to load. Would you '
    'like to try again?" The LLM is informed via a tool error result so it can respond naturally. '
    'The user can retry or continue the conversation without the app.'
)

doc.add_heading('Timeout strategy for async tool calls?', level=3)
doc.add_paragraph(
    'Default timeout: 30 seconds per tool invocation. Apps that need longer (e.g., complex computations) '
    'can declare a custom timeout in their manifest (max 120 seconds). '
    'During the wait, the UI shows a progress indicator. On timeout, the tool returns an error to the LLM: '
    '"Tool call timed out after 30s." The LLM then informs the user and suggests alternatives.'
)

doc.add_heading('How does the chatbot recover from a failed app interaction?', level=3)
doc.add_paragraph(
    'The LLM receives the error as a tool result (not a crash). It can then: '
    '(1) suggest retrying the action, (2) offer an alternative approach, or (3) continue the conversation '
    'without the app. The conversation state is never corrupted by a failed tool call because the error '
    'is recorded as a regular tool-role message in the history. '
    'The app\'s iframe can be reloaded independently without affecting the chat.'
)

doc.add_heading('Circuit breaker patterns for unreliable apps?', level=3)
doc.add_paragraph(
    'Track failure rate per app over a rolling 5-minute window. If an app fails 3 consecutive tool calls, '
    'the platform marks it as "degraded" and the LLM is informed: "The chess app is currently experiencing '
    'issues." After 5 minutes, the platform allows one retry (half-open state). If it succeeds, the app '
    'is restored; if not, it remains degraded until the next window. '
    'This prevents the LLM from repeatedly calling a broken app and frustrating the user.'
)

# --- 14. Testing Strategy ---
doc.add_heading('14. Testing Strategy', level=2)

doc.add_heading('How do you test the plugin interface in isolation?', level=3)
doc.add_paragraph(
    'Unit tests for the postMessage protocol handler using mock MessageEvent objects. '
    'Test that: (1) valid messages are processed correctly, (2) malformed messages are rejected, '
    '(3) origin validation works, (4) timeout handling fires correctly. '
    'The plugin controller is a pure TypeScript class that can be tested without a browser environment '
    'using jsdom or happy-dom (already available via Vitest).'
)

doc.add_heading('Mock apps for integration testing?', level=3)
doc.add_paragraph(
    'Create a "test-app" that implements the ChatBridge plugin protocol with configurable behavior: '
    'instant responses, delayed responses, error responses, malformed messages. '
    'This mock app runs as a simple HTML page served locally. Integration tests load it in an iframe '
    'and exercise the full tool invocation lifecycle.'
)

doc.add_heading('End-to-end testing of full invocation lifecycle?', level=3)
doc.add_paragraph(
    'Playwright for browser-based E2E tests. Test the full flow: '
    'user types "let\'s play chess" -> LLM calls chess tool -> iframe loads -> board renders -> '
    'user makes move -> app sends state update -> user asks "what should I do?" -> LLM analyzes state -> '
    'game ends -> app sends completion -> conversation continues. '
    'Mock the LLM responses to make tests deterministic.'
)

doc.add_heading('Load testing with multiple concurrent app sessions?', level=3)
doc.add_paragraph(
    'Not a priority for the MVP (single-user desktop app). For production readiness, we would use '
    'k6 or Artillery to simulate concurrent users with active app sessions, measuring: postMessage latency, '
    'LLM response time with tool calls, memory usage per iframe, and state update throughput.'
)

# --- 15. Developer Experience ---
doc.add_heading('15. Developer Experience', level=2)

doc.add_heading('How easy is it for a third-party developer to build an app?', level=3)
doc.add_paragraph(
    'Target: a developer should be able to build a basic ChatBridge app in under 2 hours. '
    'The minimum requirements are: (1) a static HTML/JS page that implements the postMessage protocol, '
    '(2) a chatbridge-manifest.json file declaring tools. '
    'We provide a JavaScript SDK (@chatbridge/plugin-sdk) that handles the postMessage boilerplate: '
    'message parsing, origin validation, type-safe event handlers, and state reporting.'
)

doc.add_heading('What documentation do they need?', level=3)
doc.add_paragraph(
    '- Quick Start guide (build your first app in 30 minutes)\n'
    '- Manifest schema reference\n'
    '- PostMessage protocol specification\n'
    '- Tool schema format (JSON Schema subset)\n'
    '- Authentication integration guide (for apps needing OAuth)\n'
    '- Example apps with source code (chess, weather, Spotify)\n'
    '- Debugging guide and common error codes'
)

doc.add_heading('Local development and testing workflow for app developers?', level=3)
doc.add_paragraph(
    'Developers run their app locally (e.g., localhost:3001) and register it with the local ChatBridge '
    'instance by pointing to their local manifest URL. The platform\'s dev mode disables origin restrictions '
    'for localhost. Hot-reload works naturally since the iframe loads from the dev server. '
    'A browser DevTools guide is provided for inspecting postMessage traffic.'
)

doc.add_heading('Debugging tools for tool invocation failures?', level=3)
doc.add_paragraph(
    'A "Plugin Debug Panel" (dev mode only) that shows: '
    '(1) all postMessage traffic (timestamped, color-coded by direction), '
    '(2) tool invocation history with parameters and results, '
    '(3) current app state summary, '
    '(4) error log with stack traces from the iframe (if the app reports them via APP_ERROR). '
    'This panel is a collapsible sidebar in the chat UI, visible only when dev mode is enabled in settings.'
)

# --- 16. Deployment & Operations ---
doc.add_heading('16. Deployment & Operations', level=2)

doc.add_heading('Where do third-party apps get hosted?', level=3)
doc.add_paragraph(
    'Third-party developers host their own apps on any static hosting (Vercel, Netlify, GitHub Pages, '
    'their own servers). The platform only stores the manifest URL and iframe entry point URL. '
    'For the 3 required demo apps, we host them on Vercel as separate projects alongside the main platform. '
    'This proves the architecture works with truly separate deployments.'
)

doc.add_heading('CI/CD for the platform itself?', level=3)
doc.add_paragraph(
    'GitHub Actions (existing in the Chatbox repo, adapted for ChatBridge): '
    'lint (Biome) -> test (Vitest) -> build (electron-vite) -> deploy (Vercel for web, '
    'electron-builder for desktop). The web version is the primary deployment target for grading. '
    'Desktop builds are secondary.'
)

doc.add_heading('Monitoring for app health and invocation success rates?', level=3)
doc.add_paragraph(
    'MVP: Console logging + Sentry (already integrated in Chatbox) for error tracking. '
    'Production: Instrument tool invocations with timing and success/failure metrics. '
    'Dashboard showing per-app: invocation count, success rate, p50/p95 latency, error breakdown. '
    'Sentry captures unhandled errors from both the platform and (if the app reports them) the iframe.'
)

doc.add_heading('How do you handle app updates without breaking existing sessions?', level=3)
doc.add_paragraph(
    'Apps are loaded fresh on each iframe mount (no caching of app code). When an app updates, '
    'the next time the iframe loads it gets the new version. Active sessions are not affected because the '
    'iframe is already loaded with the old version. The manifest includes a version field; if the platform '
    'detects a version mismatch between the cached manifest and the live one, it notifies the user: '
    '"A newer version of Chess is available. Reload?" '
    'The tool schema contract is versioned (v1) so breaking changes require a new version number.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TECHNICAL STACK DECISION SUMMARY
# ══════════════════════════════════════════════════════════════
doc.add_heading('Technical Stack Decision Summary', level=1)

table = doc.add_table(rows=11, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Layer', 'Choice', 'Rationale']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

rows_data = [
    ['Frontend', 'React 18 + Mantine + Tailwind (existing)', 'Already in the Chatbox codebase; no migration cost'],
    ['Desktop', 'Electron 26 (existing)', 'Provides iframe sandboxing and Node.js access for MCP'],
    ['Web Build', 'Vite SPA (existing web target)', 'Chatbox already supports a web build via CHATBOX_BUILD_PLATFORM=web'],
    ['State', 'Zustand + IndexedDB (existing)', 'Proven at scale in Chatbox; add a pluginStore slice'],
    ['AI/LLM', 'Vercel AI SDK + Claude Sonnet (existing)', 'Unified tool calling across providers; already integrated'],
    ['Plugin Comms', 'postMessage protocol (new)', 'Browser-native, secure, no additional dependencies'],
    ['Auth (Platform)', 'Supabase Auth (new)', 'Free tier, built-in OAuth providers, quick setup'],
    ['Auth (OAuth)', 'Supabase Auth + popup flow (new)', 'Handles Spotify OAuth; tokens stored server-side'],
    ['Database', 'Supabase PostgreSQL (new backend)', 'App registrations + OAuth tokens; free tier sufficient'],
    ['Deployment', 'Vercel (web) + Electron (desktop)', 'Vercel for grading; Electron for full feature set'],
]

for i, row_data in enumerate(rows_data):
    for j, cell_text in enumerate(row_data):
        table.rows[i + 1].cells[j].text = cell_text

doc.add_paragraph()

# ══════════════════════════════════════════════════════════════
# BUILD PRIORITY & RISK ASSESSMENT
# ══════════════════════════════════════════════════════════════
doc.add_heading('Build Priority & Risk Assessment', level=1)

doc.add_heading('Priority Order', level=2)
priorities = [
    ('1. Basic chat (Day 1)', 'Verify the Chatbox fork builds and runs. Add platform auth. Low risk - existing code.'),
    ('2. Plugin manifest + registry (Day 1-2)', 'Define the manifest schema and build the registration system. Medium risk - new code, but small surface area.'),
    ('3. PostMessage protocol (Day 2)', 'Implement the typed postMessage bridge. Medium risk - browser security edge cases.'),
    ('4. Chess app - vertical slice (Day 2-3)', 'Build chess as the first complete integration. High risk - most complex app; validates the entire architecture.'),
    ('5. Tool invocation via AI SDK (Day 3)', 'Wire plugin tools into the Vercel AI SDK streamText pipeline. Medium risk - relies on understanding the existing model-calls code.'),
    ('6. Completion signaling + context (Day 3-4)', 'App signals done, chat resumes with context. High risk - most teams struggle here per the project brief.'),
    ('7. Weather app (Day 4)', 'Public API integration, no auth. Low risk - simpler pattern.'),
    ('8. Spotify app + OAuth (Day 4-5)', 'OAuth popup flow, external API. Medium-high risk - OAuth flows have many edge cases.'),
    ('9. Error handling + polish (Day 5-6)', 'Timeouts, circuit breakers, loading states. Low risk - incremental.'),
    ('10. Documentation + deployment (Day 6-7)', 'API docs, setup guide, Vercel deploy. Low risk - time-consuming but straightforward.'),
]

for title, desc in priorities:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_heading('Top Risks', level=2)
risks = [
    ('Completion signaling complexity', 'The boundary between "app is done" and "app is still active" is ambiguous for long-running apps like chess. Mitigation: explicit APP_COMPLETE message + LLM can ask the app for status.'),
    ('Iframe security vs. functionality', 'Strict sandboxing may break apps that need cookies or localStorage. Mitigation: provide a platform-mediated storage API via postMessage.'),
    ('LLM routing accuracy', 'The LLM may invoke the wrong app or invoke an app when the user just wants to chat. Mitigation: write high-quality tool descriptions; test with ambiguous prompts.'),
    ('OAuth in iframe context', 'Popup blockers and cross-origin restrictions may interfere. Mitigation: test on Chrome, Firefox, Safari; provide fallback instructions.'),
]

for title, desc in risks:
    p = doc.add_paragraph()
    run = p.add_run(title + ': ')
    run.bold = True
    p.add_run(desc)

# ══════════════════════════════════════════════════════════════
# AI COST ANALYSIS (PRELIMINARY)
# ══════════════════════════════════════════════════════════════
doc.add_page_break()
doc.add_heading('AI Cost Analysis (Preliminary)', level=1)

doc.add_heading('Assumptions', level=2)
doc.add_paragraph(
    '- Primary model: Claude Sonnet 4 ($3/M input tokens, $15/M output tokens)\n'
    '- Routing/simple calls: Claude Haiku 4.5 ($0.80/M input, $4/M output)\n'
    '- Average conversation: 10 user messages, 10 assistant responses\n'
    '- Average tool invocations per session: 5\n'
    '- Average tokens per message: 200 input, 400 output\n'
    '- Tool schema overhead: ~1,000 tokens per call (injected tool definitions)\n'
    '- Average sessions per user per month: 20'
)

doc.add_heading('Production Cost Projections', level=2)

cost_table = doc.add_table(rows=5, cols=5)
cost_table.style = 'Light Grid Accent 1'
cost_table.alignment = WD_TABLE_ALIGNMENT.CENTER

cost_headers = ['Metric', '100 Users', '1,000 Users', '10,000 Users', '100,000 Users']
for i, h in enumerate(cost_headers):
    cost_table.rows[0].cells[i].text = h
    for p in cost_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

cost_rows = [
    ['Sessions/month', '2,000', '20,000', '200,000', '2,000,000'],
    ['Tool invocations/month', '10,000', '100,000', '1,000,000', '10,000,000'],
    ['LLM cost/month', '$45', '$450', '$4,500', '$45,000'],
    ['Infra cost/month', '$0 (free tier)', '$25', '$150', '$800'],
]

for i, row_data in enumerate(cost_rows):
    for j, cell_text in enumerate(row_data):
        cost_table.rows[i + 1].cells[j].text = cell_text

doc.add_paragraph()
doc.add_paragraph(
    'Note: At 100K users, cost optimization becomes critical. Strategies include: '
    'using Haiku for tool routing and Sonnet only for complex reasoning, caching common tool results, '
    'aggressive context compaction, and negotiating volume discounts with the LLM provider.'
)

# Save
doc.save('C:\\Users\\lramo\\Documents\\GauntletAI\\Repo\\chatbridge\\docs\\Pre-Search.docx')
print('Pre-Search.docx generated successfully.')
